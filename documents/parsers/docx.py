"""DOCX parser adapter. Requires python-docx."""
from io import BytesIO

from .base import ParsedDocument


def parse_docx(content: bytes, filename: str) -> ParsedDocument:
    from docx import Document

    document = Document(BytesIO(content))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    return ParsedDocument(
        text=text,
        metadata={"filename": filename, "format": "docx", "paragraphs": len(paragraphs)},
    )
